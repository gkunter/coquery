# -*- coding: utf-8 -*-
"""
documents.py is part of Coquery.

Copyright (c) 2016-2018 Gero Kunter (gero.kunter@coquery.org)

Coquery is released under the terms of the GNU General Public License (v3).
For details, see the file LICENSE that you should have received along
with Coquery. If not, see <http://www.gnu.org/licenses/>.
"""

from __future__ import unicode_literals

import logging
import codecs
import os.path
import re

from . import options
from .unicode import utf8

# define file type constants:


FT_BINARY = "BINARY"
FT_PLAIN = "PLAIN"
FT_HTML = "HTML"
FT_DOCX = "DOCX"
FT_ODT = "ODT"
FT_PDF = "PDF"


def detect_file_type(file_name, sample_length=1024):
    """
    Detect the file type of the supplied file.

    This function attempts to detect one of the supported text type formats:

    - PDF
    - MS Office/Word (.docx)
    - OpenDocument Text (.odt)
    - HTML

    If none of these file formats can be detected, the function tests whether
    the file is a binary or a plain text file based on the method described
    here: http://stackoverflow.com/a/7392391

    Note that this detection may give false results for some files!

    Parameters
    ----------
    file_name : str
        The path to the file.
    sample_length : int, default: 1024
        The maximum number of bytes that will be read from the beginning of
        the file.

    Returns
    -------
    file_type : str
        A string containing the file type. One of the constants defined in
        this module: FT_BINARY, FT_DOCX, FT_HTML, FT_ODT, FT_PDF, FT_PLAIN
    """

    # Define magic numbers and headers:
    _pk_header = b"PK\x03\x04"
    _pdf_header = b"%PDF-"
    _html_header = b"\s*<!DOCTYPE\s+html\s*.*>"

    # get extension:
    _, ext = os.path.splitext(file_name)

    # Read first 1024 bytes from file as a sample:
    with open(file_name, "rb") as fp:
        sample = fp.read(sample_length)

    # The following bytearray is used to detect binary files. The idea is
    # based on http://stackoverflow.com/a/7392391: all text characters (as
    # defined by the bytearray) in the sample are mapped to None. If anything
    # remains, the sample contains non-text characters, and is therefore
    # treated as a binary string.
    _textchars = bytearray({0, 7, 8, 9, 10, 12, 13, 27} |
                           set(range(0x20, 0x100)) - {0x7f})

    # try to detect file type based on the content of the sample (and,
    # where appropriate, the file extension):
    if sample[:len(_pdf_header)] == _pdf_header:
        file_type = FT_PDF
    elif sample[:len(_pk_header)] == _pk_header and ext.lower() == ".docx":
        file_type = FT_DOCX
    elif sample[:len(_pk_header)] == _pk_header and ext.lower() == ".odt":
        file_type = FT_ODT
    elif re.match(_html_header, sample, re.IGNORECASE + re.DOTALL):
        file_type = FT_HTML
    elif sample.translate(None, _textchars):
        file_type = FT_BINARY
    else:
        file_type = FT_PLAIN

    return file_type


def docx_to_str(path, encoding="utf-8"):
    if options.use_docx:
        from docx import Document

    document = Document(path)
    txt = [para.text for para in document.paragraphs]
    return "\n".join(txt)


def pdf_to_str(path, encoding="utf-8"):
    if options.use_pdfminer:
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.layout import LAParams
        from pdfminer.pdfparser import PDFParser

        # account for API change in pdfminer that didn't make it to
        # the Python 3 version:
        try:
            from pdfminer.pdfpage import PDFPage
        except ImportError:
            from pdfminer.pdfparser import PDFDocument

        # Use either cStringIO or io:
        try:
            from cStringIO import StringIO
        except ImportError:
            from io import StringIO

    content = StringIO()
    manager = PDFResourceManager()
    # not all versions of TextConverter support encodings:
    try:
        device = TextConverter(manager, content, codec=encoding,
                               laparams=LAParams())
    except TypeError:
        device = TextConverter(manager, content, laparams=LAParams())
    interpreter = PDFPageInterpreter(manager, device)

    with open(path, "rb") as pdf:
        # account for API change:
        try:
            pages = PDFPage.get_pages(pdf, set())
        except (NameError, AttributeError):
            parser = PDFParser(pdf)
            document = PDFDocument()
            parser.set_document(document)
            document.set_parser(parser)

            pages = document.get_pages()

        for page in pages:
            interpreter.process_page(page)

    txt = content.getvalue()
    content.close()

    return txt


def html_to_str(path, encoding="utf-8"):
    if options.use_bs4:
        from bs4 import BeautifulSoup

    def visible(element):
        """
        Filter from http://stackoverflow.com/a/1983219
        """
        if element.parent.name in ['style', 'script', '[document]',
                                   'head', 'title']:
            return False
        elif re.match('<!--.*-->', str(element)):
            return False
        return True

    s = plain_to_str(path)
    soup = BeautifulSoup(s, "html.parser")
    texts = soup.findAll(text=True)
    return " ".join(list(filter(visible, texts)))


def odt_to_str(path):
    if options.use_odfpy:
        from odf.opendocument import load
        from odf import text

    document = load(utf8(path))
    txt = []
    for para in document.getElementsByType(text.P):
        txt.append(utf8(para.__str__()))
    return "\n".join(txt)


def plain_to_str(path):
    if options.use_chardet:
        import chardet
        content = open(path, "rb").read()
        detection = chardet.detect(content)
        encoding = detection["encoding"]
        confidence = detection["confidence"]
        if encoding is None:
            s = ("Cannot detect encoding of file {}. Is this file really a "
                 "text file in one of the supported formats?").format(path)
            logging.warning(s)
            raw_text = content
        else:
            if confidence < 0.5:
                s = ("Low confidence ({:.2}) about the encoding of file {}. "
                     "Assuming encoding '{}'.").format(
                         confidence, path, encoding)
                logging.warning(s)
            else:
                s = "Encoding '{}' detected for file {} (confidence: {:.2})"
                s = s.format(encoding, path, confidence)
                logging.info(s)
            raw_text = content.decode(encoding)
    else:
        try:
            with codecs.open(path, "r", encoding="utf-8") as input_file:
                raw_text = input_file.read()
                s = "Assuming encoding 'utf-8' for file {}".format(path)
                logging.info(s)
        except UnicodeDecodeError:
            with codecs.open(path, "r", encoding="ISO-8859-1") as input_file:
                raw_text = input_file.read()
                s = "Assuming encoding 'ISO-8859-1' for file {}".format(path)
                logging.info(s)
    return utf8(raw_text, errors="replace")
